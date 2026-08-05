"""
Resume file parsing — PDF, DOCX, TXT.

Two failure modes worth handling explicitly rather than letting them 500:

  * A PDF that is a scan. pypdf returns almost nothing because there is no text
    layer, only pixels. Detect it and say so, instead of silently matching an
    empty resume against every job and returning nonsense scores.

  * .doc (old binary Word format). python-docx only reads .docx. Tell the person
    to re-save rather than failing with a confusing zipfile error.
"""
from __future__ import annotations

import io
import re
from typing import Tuple


class ResumeParseError(Exception):
    pass


def _clean(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ResumeParseError("PDF support is not installed. Run: pip install pypdf")

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception:
        raise ResumeParseError("That PDF could not be opened — it may be corrupted or password protected.")

    if reader.is_encrypted:
        try:
            reader.decrypt("")          # some PDFs are 'encrypted' with a blank password
        except Exception:
            raise ResumeParseError("That PDF is password protected. Remove the password and try again.")

    parts = []
    for page in reader.pages[:15]:      # a resume past 15 pages is not a resume
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    text = _clean("\n".join(parts))

    if len(text.split()) < 40:
        raise ResumeParseError(
            "Almost no text came out of that PDF. It is probably a scan or an image "
            "export — there is no text layer to read. Export it again from Word or "
            "Google Docs, or paste the text instead."
        )
    return text


def parse_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise ResumeParseError("DOCX support is not installed. Run: pip install python-docx")

    try:
        d = docx.Document(io.BytesIO(data))
    except Exception:
        raise ResumeParseError(
            "That file could not be read as a .docx. If it is an older .doc file, "
            "open it in Word and use Save As to make a .docx, then try again."
        )

    parts = [p.text for p in d.paragraphs]
    for table in d.tables:              # many resume templates are built in tables
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    text = _clean("\n".join(p for p in parts if p and p.strip()))

    if len(text.split()) < 40:
        raise ResumeParseError("That document appears to be empty.")
    return text


def parse_resume(filename: str, data: bytes) -> Tuple[str, str]:
    """Return (text, detected_kind). Raises ResumeParseError with a readable message."""
    if not data:
        raise ResumeParseError("The uploaded file was empty.")
    if len(data) > 8 * 1024 * 1024:
        raise ResumeParseError("That file is over 8MB. Resumes should be far smaller.")

    name = (filename or "").lower()

    if name.endswith(".pdf") or data[:4] == b"%PDF":
        return parse_pdf(data), "pdf"

    if name.endswith(".docx") or data[:2] == b"PK":
        return parse_docx(data), "docx"

    if name.endswith(".doc"):
        raise ResumeParseError(
            "Old .doc files are not supported. Open it in Word, Save As .docx, and upload that."
        )

    if name.endswith((".txt", ".md")) or not name:
        try:
            text = _clean(data.decode("utf-8", errors="ignore"))
        except Exception:
            raise ResumeParseError("That file could not be read as text.")
        if len(text.split()) < 40:
            raise ResumeParseError("That file has too little text to match against.")
        return text, "text"

    raise ResumeParseError("Unsupported file type. Upload a PDF, DOCX or TXT.")
