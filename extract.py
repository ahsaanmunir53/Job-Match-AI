"""
Resume text extraction.

Handles the three formats people actually have a CV in: PDF, DOCX and plain text.

A note on scanned PDFs: if someone exports their CV as an image (or scans a
printed copy), there is no text layer to pull out. Rather than returning an empty
string and letting the matcher silently score everything at zero, extraction
reports the problem so the UI can tell them to paste instead. Silent empty output
is the worse failure — it looks like the app is broken rather than the file.
"""
from __future__ import annotations

import io
import re
from typing import Dict

MAX_BYTES = 5 * 1024 * 1024          # 5MB — a CV that size is a scan, not a doc
MIN_CHARS = 120                       # below this, extraction effectively failed


def _tidy(text: str) -> str:
    """Normalise whitespace without destroying line structure."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # PDFs love to emit a bullet glyph per line; strip the noise but keep the break
    text = re.sub(r"[•▪●◦]\s*", "- ", text)
    return text.strip()


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        try:
            reader.decrypt("")            # many CVs are "encrypted" with no password
        except Exception:
            raise ValueError("This PDF is password protected. Remove the password or paste the text.")
    pages = []
    for page in reader.pages[:15]:        # a CV past 15 pages is not a CV
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(pages)


def _from_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    # Plenty of CV templates put the entire history inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract(filename: str, data: bytes) -> Dict:
    """
    Returns {text, chars, format, warning?} or raises ValueError with a message
    that is safe to show the user directly.
    """
    if not data:
        raise ValueError("That file is empty.")
    if len(data) > MAX_BYTES:
        raise ValueError("File is larger than 5MB. Export a text-based PDF or paste the text.")

    name = (filename or "").lower()

    if name.endswith(".pdf"):
        fmt, text = "PDF", _from_pdf(data)
    elif name.endswith(".docx"):
        fmt, text = "DOCX", _from_docx(data)
    elif name.endswith(".doc"):
        raise ValueError("Old .doc format is not supported. Save as .docx or PDF, or paste the text.")
    elif name.endswith((".txt", ".md", ".rtf")):
        fmt = "text"
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1", errors="ignore")
    else:
        raise ValueError("Unsupported file type. Upload a PDF, DOCX or TXT — or paste the text.")

    text = _tidy(text)

    if len(text) < MIN_CHARS:
        raise ValueError(
            f"Only {len(text)} characters could be read from that {fmt}. "
            "It is probably a scan or an image export with no text layer — "
            "please paste your resume text instead."
        )

    return {
        "text": text,
        "chars": len(text),
        "words": len(text.split()),
        "format": fmt,
        "warning": (
            "Extraction looks thin — check the text below before matching."
            if len(text.split()) < 150 else None
        ),
    }
