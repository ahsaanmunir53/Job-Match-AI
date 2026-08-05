"""
Resume parsing, skill extraction, and resume→job match scoring.

Scoring = 0.55 × (scaled TF-IDF cosine similarity between the resume and the
job text) + 0.45 × (skill-overlap ratio). Deterministic, dependency-light —
no scikit-learn, so it stays fast on a 512 MB free-tier dyno.
"""

from __future__ import annotations

import io
import math
import re
from collections import Counter

# ---------------------------------------------------------------------------
# Skill vocabulary: canonical name -> list of alias patterns (lowercase).
# Multi-word aliases are matched as phrases; single tokens on word boundaries.
# ---------------------------------------------------------------------------

SKILLS: dict[str, list[str]] = {
    # Languages
    "Python": ["python"],
    "JavaScript": ["javascript", "js"],
    "TypeScript": ["typescript", "ts"],
    "Java": ["java"],
    "C++": ["c++", "cpp"],
    "C#": ["c#", "csharp", ".net c#"],
    "Go": ["golang", "go lang"],
    "Rust": ["rust"],
    "Ruby": ["ruby"],
    "PHP": ["php"],
    "Kotlin": ["kotlin"],
    "Swift": ["swift"],
    "Dart": ["dart"],
    "Scala": ["scala"],
    "R": ["r programming"],
    "MATLAB": ["matlab"],
    "SQL": ["sql"],
    "Bash / Shell": ["bash", "shell scripting", "shell script"],
    "HTML": ["html", "html5"],
    "CSS": ["css", "css3", "scss", "sass", "tailwind", "tailwindcss", "bootstrap"],

    # Frontend
    "React": ["react", "reactjs", "react.js"],
    "Next.js": ["next.js", "nextjs"],
    "Angular": ["angular", "angularjs"],
    "Vue": ["vue", "vuejs", "vue.js", "nuxt"],
    "Svelte": ["svelte"],
    "Redux": ["redux"],
    "jQuery": ["jquery"],
    "Frontend Development": ["front-end", "frontend"],

    # Backend & frameworks
    "Node.js": ["node.js", "nodejs", "node js"],
    "Express": ["express.js", "expressjs", "express js"],
    "NestJS": ["nestjs", "nest.js"],
    "Django": ["django", "drf", "django rest framework"],
    "Flask": ["flask"],
    "FastAPI": ["fastapi", "fast api"],
    "Ruby on Rails": ["rails", "ruby on rails"],
    "Laravel": ["laravel"],
    "Spring Boot": ["spring boot", "spring framework", "springboot"],
    "ASP.NET": ["asp.net", "aspnet", ".net core", "dotnet"],
    "GraphQL": ["graphql"],
    "REST APIs": ["rest api", "rest apis", "restful", "rest services"],
    "gRPC": ["grpc"],
    "WebSockets": ["websocket", "websockets"],
    "Microservices": ["microservice", "microservices"],

    # Mobile
    "React Native": ["react native"],
    "Flutter": ["flutter"],
    "Android": ["android development", "android sdk", "android"],
    "iOS": ["ios development", "ios"],

    # Data / ML / AI
    "Machine Learning": ["machine learning", "ml engineer", "ml models"],
    "Deep Learning": ["deep learning", "neural network", "neural networks"],
    "NLP": ["nlp", "natural language processing"],
    "Computer Vision": ["computer vision", "opencv"],
    "LLM / GenAI": ["llm", "llms", "large language model", "generative ai", "genai",
                     "prompt engineering", "rag", "retrieval augmented", "langchain",
                     "llamaindex", "openai api", "anthropic", "claude api", "gpt-4",
                     "fine-tuning", "fine tuning", "bedrock", "vertex ai"],
    "TensorFlow": ["tensorflow", "keras"],
    "PyTorch": ["pytorch", "torch"],
    "scikit-learn": ["scikit-learn", "sklearn"],
    "Pandas": ["pandas"],
    "NumPy": ["numpy"],
    "Data Science": ["data science", "data scientist"],
    "Data Engineering": ["data engineering", "data engineer", "etl", "elt",
                          "data pipeline", "data pipelines"],
    "Data Analysis": ["data analysis", "data analyst", "data analytics"],
    "Power BI": ["power bi", "powerbi"],
    "Tableau": ["tableau"],
    "Apache Spark": ["spark", "pyspark"],
    "Apache Kafka": ["kafka"],
    "Airflow": ["airflow"],
    "dbt": ["dbt"],
    "Snowflake": ["snowflake"],
    "BigQuery": ["bigquery"],
    "Databricks": ["databricks"],
    "MLOps": ["mlops", "ml ops", "sagemaker", "mlflow", "kubeflow"],

    # Databases
    "PostgreSQL": ["postgres", "postgresql"],
    "MySQL": ["mysql", "mariadb"],
    "MongoDB": ["mongodb", "mongo"],
    "Redis": ["redis"],
    "Elasticsearch": ["elasticsearch", "elastic search", "opensearch"],
    "SQL Server": ["sql server", "mssql", "azure sql"],
    "Oracle DB": ["oracle database", "oracle db", "pl/sql", "plsql"],
    "DynamoDB": ["dynamodb"],
    "Cassandra": ["cassandra"],
    "SQLite": ["sqlite"],
    "Firebase": ["firebase", "firestore"],
    "Supabase": ["supabase"],
    "Vector Databases": ["pinecone", "weaviate", "qdrant", "chromadb", "pgvector",
                          "vector database", "vector db", "faiss"],

    # Cloud & DevOps
    "AWS": ["aws", "amazon web services", "ec2", "s3", "lambda", "eks", "ecs",
             "cloudformation", "route53", "cloudfront"],
    "Azure": ["azure", "aks", "azure devops", "azure functions"],
    "GCP": ["gcp", "google cloud", "gke", "cloud run"],
    "Docker": ["docker", "containerization", "containers"],
    "Kubernetes": ["kubernetes", "k8s", "helm"],
    "Terraform": ["terraform", "iac", "infrastructure as code", "pulumi"],
    "Ansible": ["ansible"],
    "CI/CD": ["ci/cd", "cicd", "continuous integration", "continuous deployment",
               "github actions", "gitlab ci", "jenkins", "circleci", "argocd",
               "argo cd"],
    "Linux": ["linux", "ubuntu", "centos", "debian"],
    "Nginx": ["nginx"],
    "DevOps": ["devops", "sre", "site reliability", "platform engineering",
                "platform engineer"],
    "Monitoring & Observability": ["prometheus", "grafana", "datadog", "new relic",
                                    "cloudwatch", "observability", "elk stack"],
    "Serverless": ["serverless"],
    "Networking": ["tcp/ip", "dns", "load balancing", "vpc", "networking"],

    # Security
    "Cybersecurity": ["cybersecurity", "cyber security", "infosec",
                       "information security", "security engineer"],
    "Penetration Testing": ["penetration testing", "pentest", "pen testing",
                             "ethical hacking", "burp suite", "kali"],
    "Application Security": ["appsec", "application security", "owasp", "sast",
                              "dast", "devsecops"],
    "SOC / SIEM": ["siem", "soc analyst", "splunk", "security operations"],
    "IAM / OAuth": ["oauth", "oauth2", "openid", "oidc", "iam", "keycloak",
                     "auth0", "jwt", "sso", "saml"],

    # QA & testing
    "QA / Testing": ["quality assurance", "qa engineer", "sqa", "test cases",
                      "manual testing", "test plans"],
    "Test Automation": ["selenium", "cypress", "playwright", "appium",
                         "automation testing", "test automation", "pytest",
                         "junit", "jest", "mocha"],
    "Load Testing": ["jmeter", "load testing", "k6", "locust"],

    # Blockchain
    "Blockchain": ["blockchain", "web3", "solidity", "smart contract",
                    "smart contracts", "ethereum", "defi", "nft"],

    # Design & product
    "UI/UX Design": ["ui/ux", "ux design", "ui design", "user experience",
                      "user interface", "wireframing", "prototyping"],
    "Figma": ["figma"],
    "Adobe Creative Suite": ["photoshop", "illustrator", "adobe xd", "after effects",
                              "premiere pro", "indesign"],
    "Graphic Design": ["graphic design", "graphic designer"],
    "Product Management": ["product manager", "product management", "product owner",
                            "roadmap", "user stories", "backlog"],
    "Project Management": ["project management", "project manager", "pmp",
                            "prince2", "gantt"],
    "Agile": ["agile", "scrum", "kanban", "sprint planning", "jira"],
    "Business Analysis": ["business analyst", "business analysis",
                           "requirements gathering", "brd", "frd"],

    # Business & ops
    "Digital Marketing": ["digital marketing", "seo", "sem", "google ads",
                           "facebook ads", "social media marketing", "ppc",
                           "content marketing", "email marketing"],
    "Sales & BD": ["business development", "b2b sales", "lead generation",
                    "sales pipeline", "crm", "hubspot", "salesforce"],
    "Customer Support": ["customer support", "customer service", "customer success",
                          "zendesk", "intercom"],
    "Accounting & Finance": ["accounting", "quickbooks", "financial analysis",
                              "bookkeeping", "audit", "acca", "cfa", "financial modeling",
                              "sap", "oracle erp", "erp"],
    "HR & Recruitment": ["recruitment", "talent acquisition", "hr operations",
                          "human resources", "payroll", "onboarding"],
    "Supply Chain": ["supply chain", "procurement", "logistics management",
                      "inventory management", "warehouse management"],
    "Content Writing": ["content writing", "copywriting", "technical writing",
                         "technical writer", "blog writing"],

    # General engineering
    "Git": ["git", "github", "gitlab", "bitbucket", "version control"],
    "OOP": ["oop", "object oriented", "object-oriented", "design patterns", "solid principles"],
    "Data Structures & Algorithms": ["data structures", "algorithms", "dsa",
                                      "problem solving", "competitive programming",
                                      "leetcode"],
    "System Design": ["system design", "distributed systems", "scalability",
                       "high availability", "architecture design"],
    "API Integration": ["api integration", "third-party apis", "payment gateway",
                         "stripe", "webhooks"],
    "WordPress": ["wordpress", "woocommerce", "shopify", "wix", "webflow"],
    "Excel": ["excel", "microsoft excel", "vba", "google sheets", "spreadsheets"],
}

# Precompile alias regexes once.
_ALIAS_PATTERNS: list[tuple[str, re.Pattern]] = []
for canonical, aliases in SKILLS.items():
    for alias in sorted(aliases, key=len, reverse=True):
        pat = re.compile(r"(?<![a-z0-9+#])" + re.escape(alias) + r"(?![a-z0-9+#])")
        _ALIAS_PATTERNS.append((canonical, pat))

_WORD_RE = re.compile(r"[a-z0-9][a-z0-9+#.\-]*")

STOPWORDS = set("""a an and are as at be by for from has have in into is it its of on or
that the to was were will with we you your our this these those they them their i my
me us if then than so not no yes do does did done can could should would may might
must about above after again against all am any because been before being below between
both but down during each few further here how more most other out over own same some
such too under until up very what when where which who whom why work job role team
company position candidate experience years skills strong ability responsibilities
requirements qualifications preferred plus etc looking join apply""".split())


# ---------------------------------------------------------------------------
# Resume text extraction
# ---------------------------------------------------------------------------

def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from PDF / DOCX / TXT bytes."""
    name = (filename or "").lower()
    if name.endswith(".pdf") or data[:5] == b"%PDF-":
        return _pdf_text(data)
    if name.endswith(".docx") or data[:2] == b"PK":
        try:
            return _docx_text(data)
        except Exception:
            pass  # fall through to plain-text attempt
    return data.decode("utf-8", errors="ignore")


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages[:20]:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def _docx_text(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Skills + tokens
# ---------------------------------------------------------------------------

def normalize(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[\u2010-\u2015]", "-", text)   # fancy dashes
    text = re.sub(r"\s+", " ", text)
    return text


def extract_skills(text: str) -> list[str]:
    """Return canonical skills present in the text, sorted alphabetically."""
    norm = normalize(text)
    found: set[str] = set()
    for canonical, pattern in _ALIAS_PATTERNS:
        if canonical in found:
            continue
        if pattern.search(norm):
            found.add(canonical)
    return sorted(found)


def tokenize(text: str) -> list[str]:
    return [t for t in _WORD_RE.findall(normalize(text))
            if t not in STOPWORDS and len(t) > 1]


# ---------------------------------------------------------------------------
# TF-IDF + cosine scoring
# ---------------------------------------------------------------------------

def _job_doc(job: dict) -> str:
    title = job.get("title") or ""
    tags = " ".join(job.get("tags") or [])
    desc = (job.get("description") or "")[:2500]
    return f"{title} {title} {title} {tags} {tags} {desc}"


def score_jobs(resume_text: str, jobs: list[dict]) -> list[dict]:
    """Attach match_score (0-100) and matched_skills to each job dict."""
    if not jobs:
        return jobs

    resume_tokens = tokenize(resume_text)
    resume_skills = set(extract_skills(resume_text))

    job_token_lists = [tokenize(_job_doc(j)) for j in jobs]

    # Document frequencies across jobs + resume
    documents = job_token_lists + [resume_tokens]
    n_docs = len(documents)
    df: Counter = Counter()
    for tokens in documents:
        df.update(set(tokens))

    def vectorize(tokens: list[str]) -> dict[str, float]:
        tf = Counter(tokens)
        total = max(1, len(tokens))
        vec = {}
        for term, count in tf.items():
            idf = math.log((n_docs + 1) / (df[term] + 1)) + 1.0
            vec[term] = (count / total) * idf
        return vec

    def cosine(a: dict[str, float], b: dict[str, float]) -> float:
        if not a or not b:
            return 0.0
        if len(b) < len(a):
            a, b = b, a
        dot = sum(w * b.get(t, 0.0) for t, w in a.items())
        na = math.sqrt(sum(w * w for w in a.values()))
        nb = math.sqrt(sum(w * w for w in b.values()))
        if na == 0 or nb == 0:
            return 0.0
        return dot / (na * nb)

    resume_vec = vectorize(resume_tokens)

    for job, tokens in zip(jobs, job_token_lists):
        job_vec = vectorize(tokens)
        cos = cosine(resume_vec, job_vec)
        cos_scaled = min(1.0, cos * 3.5)

        job_skills = set(extract_skills(_job_doc(job)))
        overlap = resume_skills & job_skills
        denom = max(3, min(len(job_skills), 12)) if job_skills else 3
        skill_ratio = min(1.0, len(overlap) / denom)

        score = round(100 * (0.55 * cos_scaled + 0.45 * skill_ratio))
        job["match_score"] = max(0, min(100, score))
        job["matched_skills"] = sorted(overlap)[:10]

    return jobs
